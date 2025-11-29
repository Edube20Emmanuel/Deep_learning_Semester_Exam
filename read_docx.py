import sys
import io

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

doc = Document('2025_UCU-CSE-EXAMS(PROJECT-BASED)-Deep-Learning-vAdvent2025-Computing&Technology (3).docx')

# Set stdout to UTF-8 encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

for para in doc.paragraphs:
    print(para.text)

# Also try reading tables if they exist
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)



